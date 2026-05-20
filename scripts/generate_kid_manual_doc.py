from pathlib import Path

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "KID_STEP_BY_STEP_MANUAL.md"
OUT = ROOT / "docs" / "KID_STEP_BY_STEP_MANUAL.docx"


def add_line(doc: Document, line: str) -> None:
    if line.startswith("# "):
        doc.add_heading(line[2:], level=1)
    elif line.startswith("## "):
        doc.add_heading(line[3:], level=2)
    elif line.startswith("### "):
        doc.add_heading(line[4:], level=3)
    elif line.startswith("- "):
        doc.add_paragraph(line[2:], style="List Bullet")
    elif line[:2].isdigit() and ". " in line[:5]:
        doc.add_paragraph(line.split(". ", 1)[1], style="List Number")
    elif line.startswith("> "):
        p = doc.add_paragraph()
        run = p.add_run(line[2:])
        run.italic = True
        run.bold = True
    elif line.strip() == "---":
        doc.add_paragraph("")
    elif line.strip():
        doc.add_paragraph(line)


def main() -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)
    for line in SOURCE.read_text(encoding="utf-8").splitlines():
        add_line(doc, line)
    doc.save(OUT)
    print(f"Generated {OUT}")


if __name__ == "__main__":
    main()

