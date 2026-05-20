from pathlib import Path

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "SHOPIFY_TIKTOK_FULL_SUMMARY.md"
OUT = ROOT / "docs" / "SHOPIFY_TIKTOK_FULL_SUMMARY.docx"


def add_markdown_line(doc: Document, line: str) -> None:
    if line.startswith("# "):
        doc.add_heading(line[2:], level=1)
    elif line.startswith("## "):
        doc.add_heading(line[3:], level=2)
    elif line.startswith("### "):
        doc.add_heading(line[4:], level=3)
    elif line.startswith("- "):
        doc.add_paragraph(line[2:], style="List Bullet")
    elif line.startswith("> "):
        p = doc.add_paragraph()
        run = p.add_run(line[2:])
        run.italic = True
    elif line.strip():
        doc.add_paragraph(line)


def main() -> None:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    for line in SOURCE.read_text(encoding="utf-8").splitlines():
        add_markdown_line(doc, line)
    doc.save(OUT)
    print(f"Generated {OUT}")


if __name__ == "__main__":
    main()

